from docx import Document

from core.knowledge.loaders.base import DocumentLoader


class DocxLoader(DocumentLoader):
    """
    DOCX document loader.
    """

    def load(self, file_path: str)->str:

        doc = Document(file_path)
        
        text_parts=[]

        for para in doc.paragraphs:
            text=para.text.strip()

            if text:
                text_parts.append(text)
        
        for table in doc.tables:
            for row in table.rows:
                row_text="\t".join(cell.text.strip() for cell in row.cells)
                if row_text:
                    text_parts.append(row_text)

        return "\n\n".join(text_parts)
    

    def supported_extension(self)->str:
        return ".docx"